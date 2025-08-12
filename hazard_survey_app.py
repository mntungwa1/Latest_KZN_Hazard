import streamlit as st
st.set_page_config(page_title="KZN Hazard Risk Assessment", layout="wide")
st.markdown("<style>div.block-container{padding-top: 1rem;}</style>", unsafe_allow_html=True)

import pandas as pd
import geopandas as gpd
import folium
try:
    from streamlit_folium import st_folium
except ImportError:
    st.error("Missing dependency: streamlit-folium. Install with 'pip install streamlit-folium'.")
    st.stop()

from shapely.geometry import Point
from pathlib import Path
from datetime import datetime, date
import os, smtplib, re, zipfile
from email.message import EmailMessage

# ----------------- SESSION DEFAULTS (for Performance panel) -----------------
if "simplify_tol" not in st.session_state:
    st.session_state["simplify_tol"] = 0.0015  # degrees; higher = faster map, less detail
if "tooltip_compact" not in st.session_state:
    st.session_state["tooltip_compact"] = False  # compact tooltip shows only UID + Ward

# ----------------- SAFE SECRET HELPER (for config use) -----------------
def _safe_secret(key):
    try:
        return st.secrets.get(key, "")
    except Exception:
        return ""

# ----------------- CONFIG -----------------
# Portable storage: uses secrets.BASE_DIR if set; otherwise:
# - Windows: C:\Temp\kzn
# - Streamlit Cloud / Linux: ./kzn_data
_base_dir_secret = _safe_secret("BASE_DIR")
BASE_DIR = Path(_base_dir_secret) if _base_dir_secret else (Path(r"C:\Temp\kzn") if os.name == "nt" else Path.cwd() / "kzn_data")

SAVE_DIR   = BASE_DIR / "Responses"
MASTER_CSV = BASE_DIR / "all_submissions.csv"

EXCEL_PATH   = Path("RiskAssessmentTool.xlsm")     # Adjust if needed
GEOJSON_PATH = Path("KZN_wards.geojson")           # Adjust if needed
LOGO_PATH = "Logo.png"
SRK_LOGO_PATH = "SRK_Logo.png"

# Ensure base dirs exist early (create folders on both Windows & Cloud)
for p in (BASE_DIR, SAVE_DIR, MASTER_CSV.parent):
    try:
        p.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass

# ----------------- HELPERS -----------------
def ensure_save_dir():
    for p in (BASE_DIR, SAVE_DIR, MASTER_CSV.parent):
        p.mkdir(parents=True, exist_ok=True)

def safe_filename(name):
    return re.sub(r'[^A-Za-z0-9_-]', '_', str(name or ""))

# Break long unspaced tokens to avoid FPDF width errors
def _safe_break(text, length=80):
    s = str(text or "")
    if not s:
        return s
    parts, token = [], ""
    for ch in s:
        token += ch
        if len(token) >= length and " " not in token:
            parts.append(token)
            token = ""
        elif ch == " ":
            parts.append(token)
            token = ""
    if token:
        parts.append(token)
    fixed = []
    for p in parts:
        if len(p) >= length and " " not in p:
            fixed.extend([p[i:i+length] for i in range(0, len(p), length)])
        else:
            fixed.append(p)
    return " ".join(fixed) if fixed else s

# ----------------- STATE RESET HELPERS -----------------
def reset_form_state():
    """Clear just the respondent/form-related keys; keep auth and performance prefs."""
    keys_to_clear = [
        "name", "district_municipality", "local_municipality",
        "final_ward", "ward_name", "today", "user_email", "extra_info",
        "hazards_selected", "custom_hazard", "selected_ward",
        "files_saved", "active_tab",
    ]
    for k in keys_to_clear:
        st.session_state.pop(k, None)

def clear_all_state():
    """Clear EVERYTHING, including login; also clear cached data."""
    st.session_state.clear()
    try:
        st.cache_data.clear()
    except Exception:
        pass
    st.rerun()

# ----------------- SECRETS (safe fallback) -----------------
def _load_secrets():
    try:
        return dict(st.secrets)
    except Exception:
        return {}

_SECRETS = _load_secrets()

def _s(key, default=""):
    if isinstance(default, str):
        return (_SECRETS.get(key, os.getenv(key, default)) or "").strip()
    val = _SECRETS.get(key, os.getenv(key))
    return val if val is not None else default

EMAIL_ADDRESS  = _s("EMAIL_ADDRESS", "")
EMAIL_PASSWORD = _s("EMAIL_PASSWORD", "")
APP_PASSWORD   = _s("APP_PASSWORD", "kzn!23@")
ADMIN_PASSWORD = _s("ADMIN_PASSWORD", "kzn!23&")
ADMIN_EMAILS   = list({_s("ADMIN_EMAIL", ""), EMAIL_ADDRESS, "dingaanm@gmail.com"} - {""})

def email_configured():
    return bool(EMAIL_ADDRESS and EMAIL_PASSWORD)

# ----------------- AUTH -----------------
def password_protection():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    password = st.text_input("Enter password to access the app:", type="password")
    if st.button("Login"):
        if password == APP_PASSWORD:
            st.session_state["authenticated"] = True
            st.success("Access granted. Please continue.")
            st.rerun()
        else:
            st.error("Incorrect password.")

# ----------------- DATA LOAD (PERFORMANCE) -----------------
@st.cache_resource(show_spinner=False)
def load_hazards_from_excel(path):
    import pandas as pd
    try:
        df = pd.read_excel(path, sheet_name="Hazard information", skiprows=1, engine="openpyxl")
    except ImportError:
        st.error("Missing dependency: openpyxl. Install with `pip install openpyxl`.")
        raise
    return df.iloc[:, 0].dropna().astype(str).tolist()

@st.cache_resource(show_spinner=False)
def load_ward_gdf(path):
    gdf = gpd.read_file(path)
    if gdf.crs is None:
        gdf.set_crs(epsg=4326, inplace=True)
    else:
        gdf = gdf.to_crs(epsg=4326)

    # Cast tooltip columns to str to avoid JSON encoding issues
    for c in ["UID", "NAMECODE", "GRID_ID", "MUNICNAME", "DISTRICT_N"]:
        if c in gdf.columns:
            try:
                gdf[c] = gdf[c].astype(str)
            except Exception:
                gdf[c] = gdf[c].apply(lambda x: "" if pd.isna(x) else str(x))

    # Spatial index: build once and keep (cache_resource persists)
    _ = gdf.sindex

    # An initial simplified geometry (will be overridden by performance panel)
    try:
        gdf["__geom_display__"] = gdf.geometry.simplify(0.0015, preserve_topology=True)
    except Exception:
        gdf["__geom_display__"] = gdf.geometry

    return gdf

def hazards_source_ui():
    if EXCEL_PATH.exists():
        try:
            return load_hazards_from_excel(EXCEL_PATH)
        except Exception as e:
            st.error("Failed to read hazards from '{}': {}".format(EXCEL_PATH, e))
    st.warning("RiskAssessmentTool.xlsm not found/unreadable. Upload an Excel with a 'Hazard information' sheet.")
    up = st.file_uploader("Upload hazards Excel", type=["xls", "xlsx", "xlsm"])
    if up:
        try:
            df = pd.read_excel(up, sheet_name="Hazard information", skiprows=1)
            return df.iloc[:, 0].dropna().astype(str).tolist()
        except Exception as e:
            st.error("Could not parse uploaded file: {}".format(e))
    return []

def wards_source_ui():
    if GEOJSON_PATH.exists():
        try:
            return load_ward_gdf(GEOJSON_PATH)
        except Exception as e:
            st.error("Failed to load wards from '{}': {}".format(GEOJSON_PATH, e))
    st.warning("KZN_wards.geojson not found/unreadable. Upload a GeoJSON of wards with a 'UID' field.")
    up = st.file_uploader("Upload wards GeoJSON", type=["geojson", "json"], key="geojson_up")
    if up:
        try:
            uploads_dir = BASE_DIR / "uploads"
            uploads_dir.mkdir(parents=True, exist_ok=True)
            tmp = uploads_dir / safe_filename(up.name)
            with open(tmp, "wb") as f:
                f.write(up.read())
            return load_ward_gdf(tmp)
        except Exception as e:
            st.error("Could not parse uploaded GeoJSON: {}".format(e))
    return None

# ----------- Field mapping (your GeoJSON) -----------
FIELD_MAP_FIXED = {
    "uid": "UID",
    "ward_name": "NAMECODE",     # auto-fill uses NAMECODE
    "ward_code": "GRID_ID",      # tooltip only
    "local_muni": "MUNICNAME",
    "district": "DISTRICT_N",
}

def get_attr(row, fld):
    try:
        return row.get(fld)
    except Exception:
        try:
            return row[fld]
        except Exception:
            return None

# ----------------- PERFORMANCE HELPERS -----------------
def apply_display_geometry(gdf, tol):
    """Return a copy of gdf with a simplified geometry column for fast map rendering."""
    gdf_out = gdf.copy()
    try:
        gdf_out["__geom_display__"] = gdf_out.geometry.simplify(float(tol), preserve_topology=True)
    except Exception:
        gdf_out["__geom_display__"] = gdf_out.geometry
    return gdf_out

def get_tooltip_fields(gdf, compact=False):
    """Build tooltip fields list based on compact toggle."""
    base = [("UID", "UID"), ("NAMECODE", "Ward")]
    extra = [("GRID_ID", "Code"), ("MUNICNAME", "Local Municipality"), ("DISTRICT_N", "District")]
    candidates = base if compact else (base + extra)

    fields, aliases, seen = [], [], set()
    for col, label in candidates:
        if col in gdf.columns and col not in seen:
            fields.append(col)
            aliases.append(label + ":")
            seen.add(col)
    return fields, aliases

# ---- NEW: Build a minimal, serializable FeatureCollection for Folium ----
def _build_geojson(gdf, fields, use_display_geom=True):
    """Return a minimal, JSON-serializable FeatureCollection for Folium."""
    feats = []
    geom_col = "__geom_display__" if (use_display_geom and "__geom_display__" in gdf.columns) else gdf.geometry.name
    for _, row in gdf.iterrows():
        geom = row[geom_col]
        if geom is None or geom.is_empty:
            continue
        props = {}
        for f in fields:
            v = row.get(f, "")
            if pd.isna(v):
                v = ""
            props[f] = str(v)
        feats.append({
            "type": "Feature",
            "properties": props,
            "geometry": geom.__geo_interface__,
        })
    return {"type": "FeatureCollection", "features": feats}

# ----------------- EMAIL -----------------
def send_email(subject, body, to_emails, attachments):
    if not email_configured():
        st.info("Email not configured (missing EMAIL_ADDRESS or EMAIL_PASSWORD). Skipping email send.")
        return
    try:
        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = EMAIL_ADDRESS
        msg["To"] = ", ".join(to_emails)
        msg.set_content(body)
        for attachment in attachments or []:
            try:
                with open(attachment, "rb") as f:
                    data = f.read()
                filename = Path(attachment).name
                msg.add_attachment(data, maintype="application", subtype="octet-stream", filename=filename)
            except Exception as e:
                st.warning("Could not attach {}: {}".format(attachment, e))
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=30) as smtp:
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)
        st.success("Email sent to {}!".format(to_emails))
    except smtplib.SMTPAuthenticationError as e:
        try:
            with smtplib.SMTP("smtp.gmail.com", 587, timeout=30) as smtp:
                smtp.ehlo()
                smtp.starttls()
                smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
                smtp.send_message(msg)
            st.success("Email sent to {}! (TLS)".format(to_emails))
        except smtplib.SMTPAuthenticationError as e2:
            st.error(
                "Gmail authentication failed. Use a Google App Password (with 2-Step Verification enabled) "
                "and ensure EMAIL_ADDRESS matches the account. Details: {}".format(e2)
            )
        except Exception as e2:
            st.error("Failed to send email: {}".format(e2))
    except Exception as e:
        st.error("Failed to send email: {}".format(e))

# ----------------- SAVE RESPONSES -----------------
def append_to_master_csv(df):
    ensure_save_dir()
    df.to_csv(MASTER_CSV, mode="a", header=not MASTER_CSV.exists(), index=False)

def save_responses(responses, name, ward_uid, email, date_filled,
                   district_municipality=None, local_municipality=None, extra_info=None, ward_name=None):
    ensure_save_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = "{}_{}_{}".format(safe_filename(ward_uid), safe_filename(name), timestamp)
    csv_path = SAVE_DIR / "{}.csv".format(base_filename)
    pdf_path = SAVE_DIR / "{}.pdf".format(base_filename)
    docx_path = SAVE_DIR / "{}.docx".format(base_filename)

    df = pd.DataFrame(responses)
    for col in ["Hazard", "Question", "Response"]:
        if col not in df.columns:
            df[col] = ""

    df.insert(0, "Respondent Name", name)
    df.insert(1, "District Municipality", district_municipality)
    df.insert(2, "Local Municipality", local_municipality)
    df.insert(3, "UID", ward_uid)
    df.insert(4, "Ward Name", ward_name)
    df.insert(5, "Email", email)
    df.insert(6, "Extra Info", extra_info)
    df.insert(7, "Date", date_filled)
    df.to_csv(csv_path, index=False)
    append_to_master_csv(df)

    # DOCX (optional, lazy import)
    try:
        from docx import Document  # lazy import so app runs even if missing
        doc = Document()
        doc.add_heading("KZN Hazard Risk Assessment Survey", 0)
        meta = [
            ("Name", name),
            ("District Municipality", district_municipality),
            ("Local Municipality", local_municipality),
            ("UID", ward_uid),
            ("Ward Name", ward_name),
            ("Email", email),
            ("Extra Info", extra_info),
            ("Date", date_filled),
        ]
        for k, v in meta:
            doc.add_paragraph("{}: {}".format(k, v))
        for _, row in df.iterrows():
            doc.add_paragraph("Hazard: {} | Question: {} | Response: {}".format(row["Hazard"], row["Question"], row["Response"]))
        doc.save(str(docx_path))
    except Exception as e:
        st.warning("DOCX not created (python-docx missing or failed): {}".format(e))
        docx_path = None

    # PDF (robust, lazy import) - FIXED for long lines & meta
    try:
        from fpdf import FPDF  # fpdf2
        pdf = FPDF()
        pdf.set_margins(12, 12, 12)
        pdf.set_auto_page_break(auto=True, margin=12)
        pdf.add_page()

        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, txt="KZN Hazard Risk Assessment Survey", ln=True, align="C")

        pdf.set_font("Arial", size=11)
        meta_lines = [
            "Name: {}".format(_safe_break(name)),
            "UID: {}".format(_safe_break(ward_uid)),
            "Ward Name: {}".format(_safe_break(ward_name)),
            "Email: {}".format(_safe_break(email)),
            "Date: {}".format(_safe_break(date_filled)),
            "District Municipality: {}".format(_safe_break(district_municipality)),
            "Local Municipality: {}".format(_safe_break(local_municipality)),
            "Extra Info: {}".format(_safe_break(extra_info)),
        ]
        for line in meta_lines:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 8, txt=str(line or ""))

        pdf.ln(2)
        pdf.set_x(pdf.l_margin)

        for _, row in df.iterrows():
            hazard   = _safe_break(row.get("Hazard", ""))
            question = _safe_break(row.get("Question", ""))
            resp     = _safe_break(row.get("Response", ""))
            text = "Hazard: {} | Question: {} | Response: {}".format(hazard, question, resp)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 7, txt=text)

        pdf.output(str(pdf_path))
    except Exception as e:
        st.warning("PDF not created (fpdf/fpdf2 missing or failed): {}".format(e))
        pdf_path = None

    return csv_path, docx_path, pdf_path

def create_zip(local_municipality, files):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_name = "{}_{}.zip".format(safe_filename(local_municipality) or "KZN", timestamp)
    zip_path = SAVE_DIR / zip_name
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zipf:
        for file in files:
            try:
                if file and Path(file).exists():
                    zipf.write(file, os.path.basename(file))
            except Exception:
                pass
    return zip_path

# ----------------- QUESTIONS -----------------
questions_with_descriptions = {
    "Has this hazard occurred in the past?": [
        "0 - Has not occurred and has no chance of occurrence",
        "1 - Has not occurred but there is real potential for occurrence",
        "2 - Has occurred but only once",
        "3 - Has occurred but only a few times or rarely",
        "4 - Has occurred regularly or at least once a year",
        "5 - Occurs multiple times during a single year",
    ],
    "How is the trend changing?": [
        "0 - Unknown / Not applicable",
        "1 - Decreasing",
        "2 - Stable",
        "3 - Marginally increasing",
        "4 - Increasing",
        "5 - Increasing rapidly",
    ],
}

capacity_questions = [
    "Sufficient staff/human resources",
    "Experience and special knowledge",
    "Equipment availability",
    "Adequate funding/budget allocation",
    "Facilities and infrastructure for response",
    "Prevention and mitigation plans",
    "Response and recovery plans",
    "Community awareness and training programs",
    "Early warning systems in place",
    "Coordination with local authorities and partners",
]

capacity_options = ["Strongly Disagree", "Disagree", "Neutral", "Agree", "Strongly Agree"]

def build_hazard_questions(hazards_to_ask):
    responses = []
    for hazard in hazards_to_ask:
        st.markdown("### {}".format(hazard))
        for q, opts in questions_with_descriptions.items():
            response = st.radio(q, opts, key="{}_{}".format(hazard, q))
            responses.append({"Hazard": hazard, "Question": q, "Response": response})
        for cq in capacity_questions:
            response = st.radio(cq, capacity_options, key="{}_{}".format(hazard, cq))
            responses.append({"Hazard": hazard, "Question": cq, "Response": response})
    return responses

# ----------------- MAP DISPLAY (PERFORMANCE) -----------------
def display_map(gdf, compact_tooltip=False):
    # Use simplified geometry if present
    gdf_display = gdf
    if "__geom_display__" in gdf.columns:
        try:
            gdf_display = gdf.set_geometry("__geom_display__", drop=False)
        except Exception:
            gdf_display = gdf

    m = folium.Map(location=[-28.6, 31.0], zoom_start=7)

    fields, aliases = get_tooltip_fields(gdf_display, compact=compact_tooltip)

    # Build cleaned, serializable GeoJSON to avoid TypeErrors in Jinja2/json
    fc = _build_geojson(gdf_display, fields, use_display_geom=True)

    gj = folium.GeoJson(
        data=fc,
        style_function=lambda x: {"fillColor": "#3186cc", "color": "black", "weight": 1, "fillOpacity": 0.35},
        highlight_function=lambda x: {"fillColor": "#ffcc00", "color": "black", "weight": 2, "fillOpacity": 0.65},
        name="wards",
    )
    if fields and (len(fields) == len(aliases)):
        gj.add_child(folium.GeoJsonTooltip(fields=fields, aliases=aliases, sticky=True))
    gj.add_to(m)
    return st_folium(m, height=700, width=1000)

def pick_feature_from_click(map_data, gdf):
    if not map_data or not map_data.get("last_clicked"):
        return None
    lng = map_data["last_clicked"].get("lng")
    lat = map_data["last_clicked"].get("lat")
    if lng is None or lat is None:
        return None
    pt = Point(lng, lat)
    try:
        sidx = gdf.sindex
        candidates = list(sidx.query(pt, predicate="intersects"))
        sub = gdf.iloc[candidates] if candidates else gdf
    except Exception:
        sub = gdf
    for _, row in sub.iterrows():
        geom = row.geometry
        if geom is not None and (geom.covers(pt) or geom.contains(pt)):
            return row
    return None

# ----------------- SURVEY -----------------
def run_survey():
    st.title("KZN Hazard Risk Assessment Survey")

    hazards = hazards_source_ui()
    gdf = wards_source_ui()
    if gdf is None:
        st.stop()

    field_map = FIELD_MAP_FIXED

    # Apply current simplification tolerance for fast display
    tol = st.session_state.get("simplify_tol", 0.0015)
    gdf_for_map = apply_display_geometry(gdf, tol)

    # Map UI + click-to-fill (compact tooltip optional)
    map_data = display_map(gdf_for_map, compact_tooltip=st.session_state.get("tooltip_compact", False))

    # Use original (unsimplified) gdf for accurate point-in-polygon
    row = pick_feature_from_click(map_data, gdf)
    if row is not None:
        clicked_uid   = get_attr(row, field_map["uid"])
        clicked_name  = get_attr(row, field_map["ward_name"])  # NAMECODE
        clicked_lm    = get_attr(row, field_map["local_muni"])
        clicked_dist  = get_attr(row, field_map["district"])

        if clicked_uid:
            st.session_state["selected_ward"] = clicked_uid
        if clicked_name:
            st.session_state["ward_name"] = clicked_name
        if clicked_lm:
            st.session_state["local_municipality"] = clicked_lm
        if clicked_dist:
            st.session_state["district_municipality"] = clicked_dist

    # Show selection confirmation
    ward_display = st.session_state.get("selected_ward", "")
    ward_name_display = st.session_state.get("ward_name", "")
    if ward_display:
        msg = "Selected UID: {}".format(ward_display)
        if ward_name_display:
            msg = "{} | Ward: {}".format(msg, ward_name_display)
        st.success(msg)

    # Hazards selection
    st.subheader("Select Applicable Hazards")
    selected = st.multiselect("Choose hazards:", hazards, key="hazards_selected")
    custom = st.text_input("Other hazard (optional)", key="custom_hazard") if st.checkbox("Add custom hazard") else ""

    if selected or custom:
        if "active_tab" not in st.session_state:
            st.session_state.active_tab = "Respondent Info"

        if st.session_state.active_tab == "Respondent Info":
            st.subheader("Respondent Info")
            st.session_state["name"] = st.text_input("Full Name", st.session_state.get("name", ""))
            st.session_state["district_municipality"] = st.text_input(
                "District Municipality",
                st.session_state.get("district_municipality", "")
            )
            st.session_state["local_municipality"] = st.text_input(
                "Local Municipality",
                st.session_state.get("local_municipality", "")
            )
            st.session_state["final_ward"] = ward_display or st.text_input(
                "UID (if not using map)",
                st.session_state.get("final_ward", "")
            )
            st.session_state["ward_name"] = st.text_input(
                "Ward Name",
                st.session_state.get("ward_name", "")
            )
            st.session_state["today"] = st.date_input("Date", value=st.session_state.get("today", date.today()))
            st.session_state["user_email"] = st.text_input("Your Email (optional for receipt)", st.session_state.get("user_email", ""))
            st.session_state["extra_info"] = st.text_area("Any extra information to be added", st.session_state.get("extra_info", ""))

            cols = st.columns(3)
            if cols[0].button("Go to Hazard Risk Evaluation"):
                st.session_state.active_tab = "Hazard Risk Evaluation"
                st.rerun()
            if cols[2].button("Clear cache (refresh data)"):
                st.cache_data.clear()
                st.rerun()

        elif st.session_state.active_tab == "Hazard Risk Evaluation":
            st.subheader("Hazard Risk Evaluation")
            hazards_to_ask = selected + ([custom] if custom else [])

            # Fast path: compact grid input (recommended)
            use_grid = st.toggle("Use compact table input (faster)", value=True, help="Edit all answers in one table")

            with st.form("hazard_form"):
                responses_to_save = None

                if use_grid:
                    rows = []
                    for hz in hazards_to_ask:
                        rows.append({
                            "Hazard": hz,
                            "Occurred": "",
                            "Trend": "",
                            "Staff": "",
                            "Experience": "",
                            "Equipment": "",
                            "Funding": "",
                            "Facilities": "",
                            "PreventionPlans": "",
                            "ResponsePlans": "",
                            "Awareness": "",
                            "EarlyWarning": "",
                            "Coordination": "",
                        })
                    df_in = pd.DataFrame(rows)

                    options_0to5 = [
                        "0 - Has not occurred and has no chance of occurrence",
                        "1 - Has not occurred but there is real potential for occurrence",
                        "2 - Has occurred but only once",
                        "3 - Has occurred but only a few times or rarely",
                        "4 - Has occurred regularly or at least once a year",
                        "5 - Occurs multiple times during a single year",
                    ]
                    trend_opts = [
                        "0 - Unknown / Not applicable",
                        "1 - Decreasing",
                        "2 - Stable",
                        "3 - Marginally increasing",
                        "4 - Increasing",
                        "5 - Increasing rapidly",
                    ]
                    likert = ["Strongly Disagree", "Disagree", "Neutral", "Agree", "Strongly Agree"]

                    cfg = {
                        "Occurred": st.column_config.SelectboxColumn(options=options_0to5, help="Has this hazard occurred in the past?"),
                        "Trend": st.column_config.SelectboxColumn(options=trend_opts, help="How is the trend changing?"),
                        "Staff": st.column_config.SelectboxColumn(options=likert),
                        "Experience": st.column_config.SelectboxColumn(options=likert),
                        "Equipment": st.column_config.SelectboxColumn(options=likert),
                        "Funding": st.column_config.SelectboxColumn(options=likert),
                        "Facilities": st.column_config.SelectboxColumn(options=likert),
                        "PreventionPlans": st.column_config.SelectboxColumn(options=likert),
                        "ResponsePlans": st.column_config.SelectboxColumn(options=likert),
                        "Awareness": st.column_config.SelectboxColumn(options=likert),
                        "EarlyWarning": st.column_config.SelectboxColumn(options=likert),
                        "Coordination": st.column_config.SelectboxColumn(options=likert),
                    }

                    edited = st.data_editor(df_in, num_rows="fixed", use_container_width=True, column_config=cfg, key="grid_editor")

                    col1, col2 = st.columns(2)
                    back = col1.form_submit_button("Back to Respondent Info")
                    submit = col2.form_submit_button("Submit Survey")

                    if back:
                        st.session_state.active_tab = "Respondent Info"
                        st.rerun()

                    if submit:
                        responses_to_save = []
                        for _, r in edited.iterrows():
                            responses_to_save.append({"Hazard": r["Hazard"], "Question": "Has this hazard occurred in the past?", "Response": r["Occurred"]})
                            responses_to_save.append({"Hazard": r["Hazard"], "Question": "How is the trend changing?", "Response": r["Trend"]})
                            for q, col in [
                                ("Sufficient staff/human resources","Staff"),
                                ("Experience and special knowledge","Experience"),
                                ("Equipment availability","Equipment"),
                                ("Adequate funding/budget allocation","Funding"),
                                ("Facilities and infrastructure for response","Facilities"),
                                ("Prevention and mitigation plans","PreventionPlans"),
                                ("Response and recovery plans","ResponsePlans"),
                                ("Community awareness and training programs","Awareness"),
                                ("Early warning systems in place","EarlyWarning"),
                                ("Coordination with local authorities and partners","Coordination"),
                            ]:
                                responses_to_save.append({"Hazard": r["Hazard"], "Question": q, "Response": r[col]})

                else:
                    responses = build_hazard_questions(hazards_to_ask)
                    col1, col2 = st.columns(2)
                    back = col1.form_submit_button("Back to Respondent Info")
                    submit = col2.form_submit_button("Submit Survey")

                    if back:
                        st.session_state.active_tab = "Respondent Info"
                        st.rerun()

                    if submit:
                        responses_to_save = responses

                if submit:
                    if not st.session_state.get("name"):
                        st.error("Please fill in your name.")
                    elif not st.session_state.get("final_ward"):
                        st.error("Please select a ward on the map or enter the UID.")
                    else:
                        csv_file, doc_file, pdf_file = save_responses(
                            responses_to_save or [],
                            st.session_state["name"],
                            st.session_state["final_ward"],
                            st.session_state["user_email"],
                            st.session_state["today"],
                            st.session_state.get("district_municipality"),
                            st.session_state.get("local_municipality"),
                            st.session_state.get("extra_info"),
                            st.session_state.get("ward_name")
                        )
                        zip_file = create_zip(st.session_state.get("local_municipality", ""), [csv_file, doc_file, pdf_file])
                        st.session_state["files_saved"] = (csv_file, doc_file, pdf_file, zip_file)
                        st.success("Survey submitted successfully! Files saved in: {}".format(SAVE_DIR))

                        if st.session_state["user_email"]:
                            send_email(
                                "Your KZN Hazard Survey Submission",
                                "Thank you for completing the survey. Your files are attached as a ZIP archive.",
                                [st.session_state["user_email"]],
                                [zip_file]
                            )
                        if ADMIN_EMAILS:
                            send_email(
                                "New KZN Hazard Survey Submission",
                                "A new survey has been submitted. See attached ZIP file.",
                                ADMIN_EMAILS,
                                [zip_file]
                            )

    if "files_saved" in st.session_state:
        csv_file, doc_file, pdf_file, zip_file = st.session_state["files_saved"]
        st.divider()
        st.caption("Downloads")
        if csv_file and Path(csv_file).exists():
            with open(csv_file, "rb") as f:
                st.download_button("Download CSV", f, file_name=os.path.basename(csv_file), mime="text/csv")
        if doc_file and (doc_file is not None) and Path(doc_file).exists():
            with open(doc_file, "rb") as f:
                st.download_button("Download DOCX", f, file_name=os.path.basename(doc_file),
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        if pdf_file and (pdf_file is not None) and Path(pdf_file).exists():
            with open(pdf_file, "rb") as f:
                st.download_button("Download PDF", f, file_name=os.path.basename(pdf_file), mime="application/pdf")
        if zip_file and Path(zip_file).exists():
            with open(zip_file, "rb") as zf:
                st.download_button("Download All (ZIP)", zf, file_name=os.path.basename(zip_file), mime="application/zip")

# ----------------- SIDEBAR & ROUTING -----------------
menu = st.sidebar.radio("Navigation", ["Survey", "Admin Dashboard"])

if os.path.exists(LOGO_PATH):
    st.sidebar.image(LOGO_PATH, width=180)
if os.path.exists(SRK_LOGO_PATH):
    st.sidebar.image(SRK_LOGO_PATH, width=160)

with st.sidebar.expander("Performance", expanded=False):
    st.caption("Tune map speed & UI payload")
    tol = st.slider(
        "Boundary simplification (degrees)",
        min_value=0.0003, max_value=0.0050,
        value=float(st.session_state.get("simplify_tol", 0.0015)),
        step=0.0001, format="%.4f",
        help="Higher = faster map (less boundary detail)"
    )
    compact = st.checkbox(
        "Compact tooltip (UID + Ward only)",
        value=bool(st.session_state.get("tooltip_compact", False))
    )
    if st.button("Apply performance settings"):
        st.session_state["simplify_tol"] = float(tol)
        st.session_state["tooltip_compact"] = bool(compact)
        st.success("Applied: tol {:.4f}, compact tooltip {}".format(tol, compact))
        st.rerun()

with st.sidebar.expander("Environment", expanded=False):
    st.caption("Runtime versions")
    import sys
    try:
        import importlib.metadata as m  # Py3.8+
    except Exception:
        import importlib_metadata as m  # fallback if needed

    def ver(pkg):
        try:
            return m.version(pkg)
        except Exception:
            return "not installed"

    st.write("Python", sys.version.split()[0])
    st.write("streamlit", ver("streamlit"))
    st.write("geopandas", ver("geopandas"))
    st.write("folium", ver("folium"))
    st.write("streamlit-folium", ver("streamlit-folium"))
    st.write("shapely", ver("shapely"))
    st.write("pyproj", ver("pyproj"))
    st.write("fiona", ver("fiona"))
    st.write("rtree", ver("rtree"))
    st.write("pandas", ver("pandas"))
    st.write("python-docx", ver("python-docx"))
    st.write("fpdf2", ver("fpdf2"))
    st.write("fpdf", ver("fpdf"))

with st.sidebar.expander("Session", expanded=False):
    st.caption("Reset form or clear app state")
    c1, c2, c3 = st.columns(3)
    if c1.button("New respondent", key="btn_reset_form"):
        reset_form_state()
        st.toast("Form cleared")
        st.rerun()
    if c2.button("Clear ALL state", key="btn_clear_all"):
        clear_all_state()
    if c3.button("Logout"):
        st.session_state.pop("authenticated", None)
        st.session_state.pop("admin_authenticated", None)
        reset_form_state()
        st.toast("Logged out")
        st.rerun()

st.sidebar.markdown(
    "<small><i>Disclaimer: The software is developed by Dingaan Mahlangu and should not be used without prior permission.</i></small>",
    unsafe_allow_html=True
)

if menu == "Survey":
    if not st.session_state.get("authenticated", False):
        st.title("KZN Hazard Risk Assessment Survey - Login")
        password_protection()
        st.stop()
    run_survey()

elif menu == "Admin Dashboard":
    st.title("Admin Dashboard - KZN Hazard Survey")
    if "admin_authenticated" not in st.session_state:
        st.session_state["admin_authenticated"] = False
    if not st.session_state["admin_authenticated"]:
        admin_password = st.text_input("Enter Admin Password:", type="password")
        if st.button("Login as Admin"):
            if admin_password == ADMIN_PASSWORD:
                st.session_state["admin_authenticated"] = True
                st.success("Admin Access Granted.")
                st.rerun()
            else:
                st.error("Incorrect Admin Password.")
        st.stop()
    ensure_save_dir()
    if MASTER_CSV.exists():
        df = pd.read_csv(MASTER_CSV)
        st.dataframe(df, use_container_width=True)
        st.download_button(
            "Download CSV",
            df.to_csv(index=False).encode("utf-8"),
            file_name="all_submissions.csv",
            mime="text/csv"
        )
    else:
        st.warning("No submissions found.")
